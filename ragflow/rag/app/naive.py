#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
from tika import parser
from io import BytesIO
from docx import Document
from timeit import default_timer as timer
import re
from deepdoc.parser.pdf_parser import PlainParser
from rag.nlp import rag_tokenizer, naive_merge, tokenize_table, tokenize_chunks, find_codec
from deepdoc.parser import PdfParser, ExcelParser, DocxParser
from rag.settings import cron_logger

class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        tbls = []
        for tb in self.doc.tables:
            html= "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i+1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, "table", html), ""))
        return [(l, "") for l in lines if l], tbls


class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None,extract_table_html=True):
        start = timer()
        callback(msg="OCR is running...")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished")
        cron_logger.info("OCR({}~{}): {}".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis finished.")
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis finished.")
        self._text_merge()
        callback(0.67, "Text merging finished")
        tbls = self._extract_table_figure(True, zoomin, extract_table_html, True)
        #self._naive_vertical_merge()
        self._concat_downward()
        #self._filter_forpages()

        cron_logger.info("layouts: {}".format(timer() - start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None,**kwargs):
    """
        Supported file formats are docx, pdf, excel, txt.
        This method apply the naive ways to chunk files.
        Successive text will be sliced into pieces using 'delimiter'.
        Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """
    chunk_token_size = kwargs.get("chunk_size", 400)
    delimiter = kwargs.get("delimiter", "\n。；！？" )
    eng = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": chunk_token_size, "delimiter": delimiter, "layout_recognize": True})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename)),
        "type": "text"
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    sections = []
    if kwargs.get("sys_file_type") == "docx":
        callback(0.1, "Start to parse.")
        sections, tbls = Docx()(filename, binary)
        res = tokenize_table(tbls, doc, eng)
        callback(0.8, "Finish parsing.")

    elif kwargs.get("sys_file_type") == "pdf":
        pdf_parser = Pdf(
        ) if parser_config.get("layout_recognize", True) else PlainParser()
        sections, tbls = pdf_parser(filename if not binary else binary,
                                    from_page=from_page, to_page=to_page, callback=callback, extract_table_html=kwargs.get('extract_table_html',True))
        res = tokenize_table(tbls, doc, eng)

    # elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
    #     callback(0.1, "Start to parse.")
    #     excel_parser = ExcelParser()
    #     sections = [(excel_parser.html(binary), "")]
    #
    # elif re.search(r"\.(txt|md)$", filename, re.IGNORECASE):
    #     callback(0.1, "Start to parse.")
    #     txt = ""
    #     if binary:
    #         encoding = find_codec(binary)
    #         txt = binary.decode(encoding)
    #     else:
    #         with open(filename, "r") as f:
    #             while True:
    #                 l = f.readline()
    #                 if not l:
    #                     break
    #                 txt += l
    #     sections = txt.split("\n")
    #     sections = [(l, "") for l in sections if l]
    #     callback(0.8, "Finish parsing.")
    #
    # elif re.search(r"\.doc$", filename, re.IGNORECASE):
    #     callback(0.1, "Start to parse.")
    #     binary = BytesIO(binary)
    #     doc_parsed = parser.from_buffer(binary)
    #     sections = doc_parsed['content'].split('\n')
    #     sections = [(l, "") for l in sections if l]
    #     callback(0.8, "Finish parsing.")

    else:
        raise NotImplementedError(
            "file type not supported yet(doc, docx, pdf, txt supported)")

    st = timer()
    chunks = naive_merge(
        sections, parser_config.get(
            "chunk_token_num", 128), parser_config.get(
            "delimiter", "\n!?。；！？"))

    res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
    cron_logger.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass


    kwargs = {"sys_file_type": "pdf"}
    result = chunk('/home/vishwastak/Downloads/Apple.pdf', callback=dummy, **kwargs)
    # result = chunk('/home/vishwastak/Downloads/Document AI Market Research.docx', callback=dummy, **kwargs)

    # chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
    print(result)